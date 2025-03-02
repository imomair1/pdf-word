import streamlit as st
import tempfile
import os
import PyMuPDF    #import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.shared import Inches
import base64
from PIL import Image
import io
import time
import json

# Set page configuration
st.set_page_config(
    page_title="PDF to Word Converter Pro",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for enhanced styling
st.markdown("""
<style>
    /* Improved CSS with better spacing and visual hierarchy */
    .main { background-color: #0F172A; color: white; }
    .main-header { 
        font-size: 2.5rem; 
        background: linear-gradient(90deg, #6366F1, #EC4899);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin: 1rem 0;
    }
    .feature-card {
        background: rgba(30, 41, 59, 0.7);
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1.5rem;
        border: 1px solid #1E3A8A;
        transition: transform 0.2s;
    }
    .feature-card:hover { transform: translateY(-3px); }
    .preview-box { 
        background: rgba(30, 41, 59, 0.9);
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .stats-badge {
        background: #1E3A8A;
        border-radius: 20px;
        padding: 0.5rem 1rem;
        margin: 0.3rem;
        display: inline-block;
    }
</style>
""", unsafe_allow_html=True)

def convert_pdf_to_docx(pdf_file, include_images=True, include_tables=True):
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_pdf.write(pdf_file.read())
    temp_pdf.close()
    
    output_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    output_docx.close()
    
    doc = Document()
    conversion_stats = {
        'pages': 0,
        'images': 0,
        'tables': 0,
        'text_blocks': 0
    }

    try:
        with pdfplumber.open(temp_pdf.name) as pdf:
            pdf_images = fitz.open(temp_pdf.name)
            conversion_stats['pages'] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages):
                # Text extraction with layout preservation
                text = page.extract_text(layout=True, x_tolerance=2, y_tolerance=2)
                if text.strip():
                    doc.add_paragraph(text)
                    conversion_stats['text_blocks'] += 1
                
                # Table extraction
                if include_tables:
                    tables = page.find_tables()
                    for table in tables:
                        data = table.extract()
                        if data:
                            conversion_stats['tables'] += 1
                            word_table = doc.add_table(rows=len(data), cols=len(data[0]))
                            for row_idx, row in enumerate(data):
                                for col_idx, cell in enumerate(row):
                                    word_table.cell(row_idx, col_idx).text = str(cell)
                
                # Image extraction
                if include_images:
                    page_images = pdf_images[page_num].get_images(full=True)
                    for img_info in page_images:
                        xref = img_info[0]
                        base_image = pdf_images.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_stream = io.BytesIO(image_bytes)
                        try:
                            img = Image.open(image_stream)
                            temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                            img.save(temp_img.name)
                            doc.add_picture(temp_img.name, width=Inches(5))
                            os.unlink(temp_img.name)
                            conversion_stats['images'] += 1
                        except Exception as e:
                            pass
                
                if page_num < len(pdf.pages) - 1:
                    doc.add_page_break()

            pdf_images.close()
        
        doc.save(output_docx.name)
        with open(output_docx.name, "rb") as f:
            docx_data = f.read()
            
    finally:
        os.unlink(temp_pdf.name)
        os.unlink(output_docx.name)
    
    return docx_data, conversion_stats

def get_download_link(data, filename):
    b64 = base64.b64encode(data).decode()
    return f'''
    <a href="data:application/octet-stream;base64,{b64}" 
       download="{filename}" 
       class="download-btn"
       style="display: inline-block;
              padding: 0.8rem 2rem;
              background: #4F46E5;
              color: white;
              border-radius: 8px;
              text-decoration: none;
              margin-top: 1rem;">
       ⬇️ Download Converted Document
    </a>
    '''

def main():
    st.markdown('<h1 class="main-header">PDF to Word Converter Pro</h1>', unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 3])
    
    with col1:
        st.markdown("### 🛠 Features")
        st.markdown("""
        <div class="feature-card">
            <h4>📊 Table Recognition</h4>
            <p>Advanced table detection with proper formatting</p>
        </div>
        <div class="feature-card">
            <h4>🖼 Image Preservation</h4>
            <p>High-quality image extraction with scaling</p>
        </div>
        <div class="feature-card">
            <h4>📜 Layout Retention</h4>
            <p>Improved text layout preservation</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### ⚙ Settings")
        include_images = st.checkbox("Include Images", value=True)
        include_tables = st.checkbox("Include Tables", value=True)
        quality = st.select_slider("Conversion Quality", 
                                 options=["Fast", "Balanced", "High Quality"],
                                 value="Balanced")
        
    with col2:
        st.markdown("### 📤 Upload PDF")
        uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
        
        if uploaded_file:
            st.markdown("### 📄 Document Preview")
            tab1, tab2 = st.tabs(["Thumbnail", "Content Preview"])
            
            with tab1:
                try:
                    with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                        page = doc.load_page(0)
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        st.image(img_data, caption="First Page Preview", width=300)
                except:
                    st.warning("Couldn't generate preview")
            
            with tab2:
                try:
                    with pdfplumber.open(uploaded_file) as pdf:
                        first_page = pdf.pages[0]
                        text = first_page.extract_text()
                        st.text_area("Extracted Text Preview", text, height=200)
                except:
                    st.warning("Couldn't extract text preview")
            
            if st.button("✨ Convert to Word", use_container_width=True):
                progress = st.progress(0)
                status = st.empty()
                
                try:
                    status.markdown("⏳ Starting conversion...")
                    progress.progress(10)
                    time.sleep(0.5)
                    
                    status.markdown("🔍 Analyzing document structure...")
                    progress.progress(30)
                    time.sleep(0.5)
                    
                    status.markdown("📝 Extracting text and tables...")
                    docx_data, stats = convert_pdf_to_docx(
                        uploaded_file, 
                        include_images=include_images,
                        include_tables=include_tables
                    )
                    progress.progress(80)
                    
                    status.markdown("🎉 Finalizing document...")
                    time.sleep(0.5)
                    progress.progress(100)
                    
                    st.markdown("### ✅ Conversion Complete")
                    st.markdown(f'''
                    <div class="preview-box">
                        <h4>Document Statistics</h4>
                        <div class="stats-badge">📄 Pages: {stats['pages']}</div>
                        <div class="stats-badge">🖼 Images: {stats['images']}</div>
                        <div class="stats-badge">📊 Tables: {stats['tables']}</div>
                        <div class="stats-badge">📝 Text Blocks: {stats['text_blocks']}</div>
                    </div>
                    ''', unsafe_allow_html=True)
                    
                    filename = uploaded_file.name.replace(".pdf", "_converted.docx")
                    st.markdown(get_download_link(docx_data, filename), unsafe_allow_html=True)
                    
                except Exception as e:
                    progress.empty()
                    status.error(f"Conversion failed: {str(e)}")
                    st.error("""
                    🛑 Conversion Error
                    - Check if the PDF is password protected
                    - Ensure the file isn't corrupted
                    - Try a simpler document first
                    """)

if __name__ == "__main__":
    main()
